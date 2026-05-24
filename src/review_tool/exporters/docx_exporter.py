"""DOCX 审稿报告导出。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..schemas import ReviewRunResult

_BUCKET_LABELS = {
    "must_fix": "三、必须修改",
    "should_fix": "四、建议修改",
    "optional": "五、可选优化",
    "reference": "六、仅供参考",
}

_RISK_LABELS = {"high": "高", "medium": "中", "low": "低", "info": "信息"}
_CONF_LABELS = {"high": "高", "medium": "中", "low": "低"}


def export_docx(result: ReviewRunResult, output_path: Path) -> None:
    """将审稿结果导出为 DOCX 报告文件。"""
    report = result.final_report or {}
    summary = report.get("summary", {})
    selector = report.get("selector", {})
    buckets = report.get("buckets", {})
    agent_exec = report.get("agent_execution", [])
    knowledge = report.get("knowledge_used", [])
    warnings = report.get("warnings", [])
    errors = report.get("errors", [])

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_default_font(doc)

    title = doc.add_heading("智能审稿报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、审稿总览", level=1)
    _add_kv(doc, "总体风险等级", _label(summary.get("overall_risk_level", "unknown")))
    _add_kv(doc, "必须修改", f"{summary.get('must_fix_count', 0)} 条")
    _add_kv(doc, "建议修改", f"{summary.get('should_fix_count', 0)} 条")
    _add_kv(doc, "可选优化", f"{summary.get('optional_count', 0)} 条")
    _add_kv(doc, "仅供参考", f"{summary.get('reference_count', 0)} 条")
    _add_kv(
        doc,
        "启用 Agent",
        ", ".join(result.selected_agent_ids) if result.selected_agent_ids else "无",
    )
    if result.article.event_background:
        _add_kv(doc, "事件背景", result.article.event_background)

    doc.add_heading("二、Selector Agent 判断", level=1)
    _add_kv(doc, "判断稿件类型", selector.get("detected_article_type", "unknown"))
    _add_kv(doc, "识别标签", ", ".join(selector.get("detected_tags", [])) or "无")
    _add_kv(doc, "选择理由摘要", selector.get("reasoning_summary", "无"))

    for bucket_key in ["must_fix", "should_fix", "optional", "reference"]:
        doc.add_heading(_BUCKET_LABELS[bucket_key], level=1)
        items = buckets.get(bucket_key, [])
        if not items:
            doc.add_paragraph("无")
            continue

        for i, item in enumerate(items, 1):
            doc.add_heading(f"{i}. {item.get('issue_type', '未分类')}", level=2)
            _add_kv(
                doc,
                "风险等级",
                _RISK_LABELS.get(item.get("risk_level", "info"), item.get("risk_level", "")),
            )
            _add_kv(
                doc,
                "置信度",
                _CONF_LABELS.get(item.get("confidence", "low"), item.get("confidence", "")),
            )
            _add_kv(doc, "来源 Agent", item.get("agent_name", item.get("agent_id", "")))
            _add_kv(doc, "原文摘录", item.get("original_quote", ""))
            _add_kv(doc, "问题", item.get("problem", ""))
            _add_kv(doc, "可能后果", item.get("possible_consequence", ""))
            _add_kv(doc, "修改建议", item.get("suggestion", ""))

            evidence = item.get("evidence", [])
            if evidence:
                p = doc.add_paragraph()
                p.add_run("依据：").bold = True
                for ev in evidence:
                    src = ev.get("source_type", "")
                    quote = ev.get("quote", "")
                    note = ev.get("note")
                    suffix = f"（{note}）" if note else ""
                    doc.add_paragraph(f"[{src}] {quote}{suffix}", style="List Bullet")

    doc.add_heading("七、Agent 执行明细", level=1)
    if agent_exec:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Agent"
        hdr[1].text = "ID"
        hdr[2].text = "状态"
        hdr[3].text = "发现数"
        hdr[4].text = "错误"
        for ae in agent_exec:
            row = table.add_row().cells
            row[0].text = str(ae.get("agent_name", ""))
            row[1].text = str(ae.get("agent_id", ""))
            row[2].text = str(ae.get("status", ""))
            row[3].text = str(ae.get("findings_count", 0))
            row[4].text = str(ae.get("error") or "")
    else:
        doc.add_paragraph("无 Agent 执行记录。")

    doc.add_heading("八、知识库引用摘要", level=1)
    if knowledge:
        for k in knowledge:
            doc.add_paragraph(
                f"[{k.get('source_type', '')}] {k.get('title', '')} ({k.get('doc_id', '')})",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("本次未使用本地知识库材料。")

    if warnings:
        doc.add_heading("九、警告", level=1)
        for warning in warnings:
            doc.add_paragraph(str(warning), style="List Bullet")

    if errors:
        doc.add_heading("十、错误", level=1)
        for error in errors:
            if isinstance(error, dict):
                doc.add_paragraph(
                    f"[{error.get('stage', '')}] {error.get('message', '')}",
                    style="List Bullet",
                )
            else:
                doc.add_paragraph(str(error), style="List Bullet")

    doc.save(output_path)


def _add_kv(doc: Document, key: str, value: object) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{key}：").bold = True
    p.add_run(str(value) if value is not None else "")


def _label(text: str) -> str:
    labels = {"high": "高", "medium": "中", "low": "低", "info": "信息"}
    return labels.get(text, text)


def _set_default_font(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
